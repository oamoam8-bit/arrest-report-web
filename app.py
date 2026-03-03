import streamlit as st
import google.generativeai as genai
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# ==========================================
# 🔑 ส่วนตั้งค่าระบบ & ปรับแต่ง UI (CSS)
# ==========================================
API_KEY = st.secrets["GEMINI_API_KEY"]

st.set_page_config(page_title="e-Arrest Report System", layout="wide", page_icon="🛡️")

# ฝัง CSS เพื่อปรับให้หน้าตาดูเป็นทางการ (Navy Blue Tone / Clean UI)
st.markdown("""
<style>
    /* ปรับสีตัวอักษรหัวข้อหลักให้เป็นสีกรมท่า */
    h1, h2, h3 {
        color: #0F2C59;
        font-family: 'Sarabun', 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    /* ปรับแต่งกรอบ Container ให้ดูเป็นแฟ้มเอกสาร */
    div[data-testid="stContainer"] {
        border: 1px solid #D6DCE5;
        border-radius: 4px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        background-color: #FFFFFF;
        padding: 10px;
    }
    /* ปรับแต่งปุ่มกดปกติ */
    div.stButton > button {
        border-radius: 4px;
        border: 1px solid #0F2C59;
        color: #0F2C59;
        font-weight: bold;
    }
    /* ปรับแต่งปุ่มกด Primary (ปุ่มหลัก) */
    div.stButton > button[data-baseweb="button"]:hover {
        background-color: #0F2C59;
        color: #FFFFFF;
        border: 1px solid #0F2C59;
    }
    /* ปรับฉากหลังของหน้าเว็บเล็กน้อยให้ดูสะอาดตา */
    .stApp {
        background-color: #F8F9FA;
    }
</style>
""", unsafe_allow_html=True)

st.title("ระบบจัดทำบันทึกการจับกุมอิเล็กทรอนิกส์ (e-Arrest Report)")
st.markdown("**ศูนย์ปฏิบัติการ งานสืบสวน | ระบบสนับสนุนการจัดทำสำนวนการจับกุมและประมวลผลข้อกฎหมาย**")
st.markdown("---")

# --- Initialize Session State สำหรับผู้ต้องหาและของกลาง ---
if 'suspects' not in st.session_state:
    st.session_state.suspects = [{'name': '', 'id': '', 'nationality': 'ไทย', 'age': '', 'address': ''}]
if 'evidences' not in st.session_state:
    st.session_state.evidences = [{'detail': '', 'location': ''}]
if 'show_preview' not in st.session_state:
    st.session_state.show_preview = False

# ==========================================
# 📝 ส่วนที่ 1: ข้อมูลทั่วไป
# ==========================================
with st.container():
    st.markdown("### ส่วนที่ 1: ข้อมูลการจับกุมพื้นฐาน")
    
    col1, col2 = st.columns(2)
    with col1:
        report_loc = st.text_input("สถานที่ทำบันทึก (เพื่อใช้ในประโยคท้าย)", placeholder="เช่น กก.สส.ภ.จว.ภูเก็ต")
        report_date = st.text_input("วัน/เดือน/ปี และเวลา ที่บันทึก", placeholder="เช่น วันที่ 23 กันยายน 2568 เวลา 15.30 น.")
        arrest_date = st.text_input("วัน/เดือน/ปี และเวลา ที่จับกุม", placeholder="เช่น วันที่ 23 กันยายน 2568 เวลา 13.20 น.")
        arrest_loc = st.text_area("สถานที่เกิดเหตุ/จับกุม", height=100)
    
    with col2:
        commanders = st.text_area("อำนวยการจับกุมโดย", placeholder="พ.ต.อ... ผกก..., พ.ต.ท... รอง ผกก...", height=68)
        officers = st.text_area("เจ้าหน้าที่ตำรวจชุดจับกุม ได้แก่", placeholder="พิมพ์ชื่อคั่นด้วยลูกน้ำ เช่น: พ.ต.ท.สมชาย, ร.ต.อ.รักชาติ, ด.ต.ยอดเยี่ยม", height=100)
        
        st.markdown("**ข้อมูลการแจ้ง พ.ร.บ.ป้องกันและปราบปรามการทรมานฯ**")
        notify_date = st.text_input("วันที่แจ้งข้อมูล", placeholder="เช่น 23 ก.ย. 2568")
        col_n1, col_n2 = st.columns(2)
        with col_n1:
            notify_attorney = st.text_input("เวลาแจ้งอัยการ", placeholder="เช่น 14.00 น.")
        with col_n2:
            notify_district = st.text_input("เวลาแจ้งนายอำเภอ", placeholder="เช่น 14.15 น.")

# ==========================================
# 👤 ส่วนที่ 2: ข้อมูลผู้ต้องหา (เพิ่มได้หลายคน)
# ==========================================
with st.container():
    st.markdown("### ส่วนที่ 2: ข้อมูลผู้ต้องหา")
    
    for i in range(len(st.session_state.suspects)):
        st.markdown(f"**ผู้ต้องหารายที่ {i+1}**")
        c1, c2, c3 = st.columns([2, 2, 1])
        st.session_state.suspects[i]['name'] = c1.text_input("ชื่อ-นามสกุล", value=st.session_state.suspects[i]['name'], key=f"s_name_{i}")
        st.session_state.suspects[i]['id'] = c2.text_input("เลขประจำตัวประชาชน/พาสปอร์ต", value=st.session_state.suspects[i]['id'], key=f"s_id_{i}")
        st.session_state.suspects[i]['age'] = c3.text_input("อายุ", value=st.session_state.suspects[i]['age'], key=f"s_age_{i}")
        
        c4, c5 = st.columns([1, 4])
        st.session_state.suspects[i]['nationality'] = c4.text_input("สัญชาติ", value=st.session_state.suspects[i]['nationality'], key=f"s_nat_{i}")
        st.session_state.suspects[i]['address'] = c5.text_input("ภูมิลำเนา/ที่อยู่", value=st.session_state.suspects[i]['address'], key=f"s_add_{i}")
        st.markdown("---")
        
    col_btn_s1, col_btn_s2 = st.columns([1, 5])
    with col_btn_s1:
        if st.button("+ เพิ่มรายชื่อผู้ต้องหา", use_container_width=True):
            st.session_state.suspects.append({'name': '', 'id': '', 'nationality': 'ไทย', 'age': '', 'address': ''})
            st.rerun()
    with col_btn_s2:
        if len(st.session_state.suspects) > 1:
            if st.button("- ลบรายการล่าสุด", key="del_s"):
                st.session_state.suspects.pop()
                st.rerun()

# ==========================================
# 📦 ส่วนที่ 3: ข้อมูลของกลาง (เพิ่มได้หลายรายการ)
# ==========================================
with st.container():
    st.markdown("### ส่วนที่ 3: บัญชีของกลาง")
    
    for i in range(len(st.session_state.evidences)):
        col_e1, col_e2 = st.columns([2, 1])
        st.session_state.evidences[i]['detail'] = col_e1.text_input(f"รายการของกลางลำดับที่ {i+1}", value=st.session_state.evidences[i]['detail'], placeholder="ระบุรายละเอียดของกลาง", key=f"e_det_{i}")
        st.session_state.evidences[i]['location'] = col_e2.text_input(f"สถานที่/จุดที่ตรวจพบ (รายการที่ {i+1})", value=st.session_state.evidences[i]['location'], placeholder="ระบุจุดตรวจยึด", key=f"e_loc_{i}")
        
    col_btn_e1, col_btn_e2 = st.columns([1, 5])
    with col_btn_e1:
        if st.button("+ เพิ่มรายการของกลาง", use_container_width=True):
            st.session_state.evidences.append({'detail': '', 'location': ''})
            st.rerun()
    with col_btn_e2:
        if len(st.session_state.evidences) > 1:
            if st.button("- ลบรายการล่าสุด", key="del_e"):
                st.session_state.evidences.pop()
                st.rerun()

# ==========================================
# ⚖️ ส่วนที่ 4: ข้อหาและพฤติการณ์
# ==========================================
with st.container():
    st.markdown("### ส่วนที่ 4: ฐานความผิดและพฤติการณ์แห่งการจับกุม")
    
    charge_input = st.text_area("ข้อกล่าวหา / ฐานความผิด", placeholder="ระบุข้อกล่าวหา...", height=68)
    behavior_input = st.text_area("พฤติการณ์แห่งการจับกุม (ฉบับสมบูรณ์)", height=200, placeholder="ระบุพฤติการณ์ หรือคัดลอกข้อความจากระบบประมวลผลอัจฉริยะ (AI) มาวางที่นี่...")
    suspect_statement = st.text_area("คำให้การของผู้ถูกจับในชั้นจับกุม", placeholder="ระบุคำให้การเบื้องต้น เช่น ในชั้นจับกุม ผู้ถูกจับให้การรับสารภาพตลอดข้อกล่าวหา โดยให้การว่า...", height=68)

# --- จัดเตรียมข้อมูลให้อยู่ในรูปแบบข้อความเพื่อส่งให้ AI และ Preview ---
suspect_text_ai = "\n".join([f"{i+1}. {s['name']} อายุ {s['age']} ปี สัญชาติ {s['nationality']} เลขประจำตัว {s['id']} ที่อยู่ {s['address']}" for i, s in enumerate(st.session_state.suspects) if s['name']])
evidence_text_ai = "\n".join([f"{i+1}. {e['detail']} (พบที่: {e['location']})" for i, e in enumerate(st.session_state.evidences) if e['detail']])
has_evidence = len([e for e in st.session_state.evidences if e['detail']]) > 0
evidence_ending_phrase = "พร้อมด้วยของกลาง " if has_evidence else ""
ending_sentence = f"เจ้าพนักงานตำรวจชุดจับกุมจึงได้ควบคุมตัวผู้ต้องหา {evidence_ending_phrase}มาทำบันทึกจับกุมที่ {report_loc} จากนั้นนำตัวผู้ต้องหา {evidence_ending_phrase}ส่งพนักงานสอบสวนเพื่อดำเนินคดีตามกฎหมาย"

# ==========================================
# 🤖 ส่วนที่ 5: ผู้ช่วย AI
# ==========================================
with st.container():
    st.markdown("### ส่วนที่ 5: ระบบประมวลผลอัจฉริยะ (AI Assistant)")
    st.info("คำแนะนำ: เมื่อระบบประมวลผลข้อความเสร็จสิ้น กรุณาคัดลอกข้อความที่ได้ไปวางใน 'ส่วนที่ 4' ด้านบน")

    st.markdown("**5.1 สร้างพฤติการณ์อัตโนมัติ (กรณีไม่มีข้อมูลพฤติการณ์เบื้องต้น)**")
    if st.button("ประมวลผลร่างพฤติการณ์จากข้อมูลทั้งหมด (Generate Report)"):
        if not charge_input or not arrest_loc or not suspect_text_ai:
            st.error("กรุณากรอกข้อมูล 'ผู้ต้องหา', 'สถานที่เกิดเหตุ', และ 'ข้อกล่าวหา' ให้ครบถ้วนก่อนทำการประมวลผล")
        else:
            try:
                genai.configure(api_key=API_KEY)
                model = genai.GenerativeModel('gemini-2.5-flash')
                prompt = (
                    f"คุณคือพนักงานสืบสวน จงเขียน 'พฤติการณ์การจับกุม' ฉบับสมบูรณ์ \n"
                    f"ข้อห้ามเด็ดขาด: \n"
                    f"1. ห้ามเขียนแบบรายงานนาย (ห้ามใช้ เรียน, เรื่อง) ให้ใช้คำว่า 'เจ้าพนักงานตำรวจชุดจับกุม'\n"
                    f"2. ประโยคสุดท้าย **บังคับต้องจบด้วยคำนี้เป๊ะๆห้ามเปลี่ยน**: '{ending_sentence}'\n\n"
                    f"ข้อมูลสำหรับแต่งเรื่อง:\n"
                    f"- วันเวลา: {arrest_date}\n"
                    f"- สถานที่: {arrest_loc}\n"
                    f"- ชุดจับกุม: {officers}\n"
                    f"- ผู้ต้องหา:\n{suspect_text_ai}\n"
                    f"- ของกลาง:\n{evidence_text_ai}\n"
                    f"- ข้อหา: {charge_input}\n"
                    f"- คำให้การ: {suspect_statement}\n"
                )
                with st.spinner('กำลังประมวลผลร่างพฤติการณ์...'):
                    st.session_state['ai_result_1'] = model.generate_content(prompt).text
            except Exception as e:
                st.error(f"ระบบขัดข้อง: {e}")
    
    if 'ai_result_1' in st.session_state:
        st.success("ข้อความที่ระบบสร้างขึ้น:")
        st.write(st.session_state['ai_result_1'])

    st.markdown("---")

    # 🔹 อัปเดตส่วน 5.2 เน้นย้ำว่าวิเคราะห์จากพฤติการณ์ที่กรอกเท่านั้น
    st.markdown("**5.2 วิเคราะห์ข้อกฎหมายและฐานความผิด (จากพฤติการณ์ที่ได้กรอก)**")
    if st.button("วิเคราะห์ข้อกฎหมายและมาตรา (Analyze Charges)"):
        if not behavior_input:
            st.error("กรุณาระบุพฤติการณ์ใน 'ส่วนที่ 4' ก่อนทำการวิเคราะห์")
        else:
            try:
                genai.configure(api_key=API_KEY)
                model = genai.GenerativeModel('gemini-2.5-flash')
                prompt = (
                    f"วิเคราะห์ฐานความผิด 'จากพฤติการณ์ที่ได้กรอก' ดังต่อไปนี้:\n"
                    f"พฤติการณ์: '{behavior_input}'\n"
                    f"ของกลาง: '{evidence_text_ai}'\n\n"
                    f"จงระบุ 'ข้อกล่าวหา' พร้อม 'ชื่อพระราชบัญญัติ และ มาตราที่เกี่ยวข้อง' ที่ถูกต้องตามกฎหมายไทย\n"
                    f"คำสั่งบังคับเด็ดขาด:\n"
                    f"1. ให้อิงจากพฤติการณ์และข้อเท็จจริงที่ผู้ใช้งานกรอกมาให้เท่านั้น ห้ามแต่งเติมข้อเท็จจริงหรือทึกทักเอาเอง\n"
                    f"2. ต้องใช้กฎหมายไทยที่อัปเดตบังคับใช้ล่าสุดเท่านั้น ห้ามดึงกฎหมายที่ถูกยกเลิกมาตอบเด็ดขาด\n"
                    f"3. ตอบแยกเป็นข้อๆ ให้ชัดเจน เหมาะสำหรับนำไปใช้แจ้งข้อกล่าวหาในบันทึกจับกุม"
                )
                with st.spinner('กำลังสืบค้นข้อกฎหมาย...'):
                    st.session_state['ai_result_2'] = model.generate_content(prompt).text
            except Exception as e:
                st.error(f"ระบบขัดข้อง: {e}")
                
    if 'ai_result_2' in st.session_state:
        st.success("ผลการวิเคราะห์ข้อกฎหมาย (อิงจากพฤติการณ์ที่ได้กรอก):")
        st.write(st.session_state['ai_result_2'])

    st.markdown("---")

    st.markdown("**5.3 ตรวจทานและขยายความพฤติการณ์ให้สมบูรณ์ (Refine Report)**")
    if st.button("ตรวจทานและเรียบเรียงพฤติการณ์ (Refine & Expand)"):
        if not behavior_input:
            st.error("กรุณาระบุพฤติการณ์ฉบับร่างใน 'ส่วนที่ 4' ก่อนทำการตรวจทาน")
        else:
            try:
                genai.configure(api_key=API_KEY)
                model = genai.GenerativeModel('gemini-2.5-flash')
                prompt = (
                    f"นำเหตุการณ์ย่อนี้: '{behavior_input}'\n"
                    f"มาเกลาเป็น 'พฤติการณ์การจับกุม' ฉบับเต็ม แทรกข้อมูลเหล่านี้ลงไปให้ครบ:\n"
                    f"ผู้ต้องหา:\n{suspect_text_ai}\n"
                    f"ของกลาง:\n{evidence_text_ai}\n\n"
                    f"ข้อห้ามเด็ดขาด: ประโยคสุดท้ายของพฤติการณ์ **บังคับต้องจบด้วยคำนี้เป๊ะๆห้ามเปลี่ยน**: '{ending_sentence}'\n"
                    f"ให้เขียนด้วยภาษากฎหมายสละสลวย"
                )
                with st.spinner('กำลังเรียบเรียงพฤติการณ์...'):
                    st.session_state['ai_result_3'] = model.generate_content(prompt).text
            except Exception as e:
                st.error(f"ระบบขัดข้อง: {e}")
                
    if 'ai_result_3' in st.session_state:
        st.success("ข้อความที่ผ่านการเรียบเรียงแล้ว:")
        st.write(st.session_state['ai_result_3'])

# ==========================================
# 📄 ส่วนที่ 6: สร้างไฟล์ Word และดูตัวอย่าง
# ==========================================
st.markdown("### ส่วนที่ 6: ตรวจทานและส่งออกเอกสาร")

col_preview, col_export = st.columns(2)

with col_preview:
    if st.button("แสดงตัวอย่างบันทึกจับกุม (Preview Document)", use_container_width=True):
        st.session_state.show_preview = True

with col_export:
    def create_word_doc():
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'TH Sarabun PSK'
        font.size = Pt(16)
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run("บันทึกการจับกุม").bold = True
        
        doc.add_paragraph(f"สถานที่ทำบันทึก\t\t{report_loc}")
        doc.add_paragraph(f"วัน/เดือน/ปี ที่บันทึก\t{report_date}")
        doc.add_paragraph(f"วัน/เดือน/ปี ที่จับกุม\t{arrest_date}")
        doc.add_paragraph(f"สถานที่เกิดเหตุ/จับกุม\t{arrest_loc}\n")
        
        p_cmd = doc.add_paragraph()
        p_cmd.add_run("นามเจ้าพนักงานผู้จับภายใต้การอำนวยการของ ").bold = True
        p_cmd.add_run(f"{commanders}\n")
        p_cmd.add_run("เจ้าหน้าที่ตำรวจชุดจับกุม ได้แก่ ").bold = True
        p_cmd.add_run(f"{officers}\n")
        
        # --- พิมพ์รายชื่อผู้ต้องหาทั้งหมด ---
        p_sus = doc.add_paragraph()
        p_sus.add_run("ได้ร่วมกันจับกุมตัวผู้ต้องหา คือ\n").bold = True
        for i, s in enumerate(st.session_state.suspects):
            if s['name']:
                p_sus.add_run(f"{i+1}. {s['name']} อายุ {s['age']} ปี สัญชาติ {s['nationality']} เลขประจำตัว {s['id']} ที่อยู่ {s['address']}\n")
        
        # --- พิมพ์ของกลางทั้งหมด ---
        doc.add_paragraph().add_run("พร้อมด้วยของกลาง").bold = True
        p_ev = doc.add_paragraph()
        if has_evidence:
            for i, e in enumerate(st.session_state.evidences):
                if e['detail']:
                    p_ev.add_run(f"{i+1}. {e['detail']} (พบที่: {e['location']})\n")
        else:
            p_ev.add_run("- ไม่มีของกลาง -\n")
        
        p_charge = doc.add_paragraph()
        p_charge.add_run("โดยกล่าวหาว่า ").bold = True
        p_charge.add_run(f"“{charge_input}”\n")
        
        doc.add_paragraph().add_run("พร้อมได้แจ้งสิทธิของผู้ถูกจับให้ทราบถึงสิทธิตามกฎหมายตั้งแต่โอกาสแรกที่ถูกจับกุมแล้ว ดังนี้").bold = True
        doc.add_paragraph("1. มีสิทธิที่จะให้การหรือไม่ให้การก็ได้ และถ้อยคำของผู้ถูกจับอาจใช้เป็นพยานหลักฐานในการพิจารณาคดีได้")
        doc.add_paragraph("2. มีสิทธิที่จะพบและปรึกษาทนายความเป็นการเฉพาะตัว")
        doc.add_paragraph("3. มีสิทธิแจ้งให้ญาติหรือผู้ซึ่งตนไว้วางใจทราบถึงการจับกุม (ถ้าไม่เป็นอุปสรรคต่อการจับกุม หรือควบคุม และ/หรือปัญหาด้านความปลอดภัย)\n"
                          "ผู้ถูกจับได้รับทราบและเข้าใจถึงวัตถุประสงค์และเงื่อนไขของกฎหมายข้างต้นดีแล้ว\n")
        
        p_beh = doc.add_paragraph()
        p_beh.add_run("พฤติการณ์ในการจับกุม กล่าวคือ ").bold = True
        p_beh.add_run(f"{behavior_input}\n")
        
        p_state = doc.add_paragraph()
        p_state.add_run("ในชั้นจับกุม ").bold = True
        p_state.add_run(f"{suspect_statement}\n")
        
        doc.add_paragraph("ในการจับครั้งนี้ เจ้าพนักงานผู้จับทุกนายได้ปฏิบัติตามอำนาจหน้าที่ตามกฎหมาย มิได้บังคับ ขู่เข็ญ หลอกลวง ทำร้ายร่างกาย "
                          "หรือทำอันตรายแก่กาย หรือจิตใจผู้ใด หรือให้สัญญาอื่นใดที่กระทำโดยมิชอบกับผู้ต้องหาแต่อย่างใด "
                          "และมิได้ทำให้ทรัพย์สินของผู้ใด สูญหาย เสียหาย หรือไร้ประโยชน์แต่อย่างใด")
        
        doc.add_paragraph("ในการควบคุมตัวผู้ถูกจับ เจ้าหน้าที่ผู้จับกุมได้ทำการบันทึกภาพและเสียงอย่างต่อเนื่องในขณะจับและควบคุมตัวผู้ถูกจับ"
                          "ในชั้นจับกุมจนกระทั่งส่งตัวให้พนักงานสอบสวน ตามมาตรา ๒๒ วรรคหนึ่ง แห่ง พ.ร.บ.ป้องกันและปราบปรามการทรมาน"
                          "และการกระทำให้สูญหาย พ.ศ.๒๕๖๕")
        
        doc.add_paragraph("ผู้จับกุมไม่ได้กระทำการใดๆ อันเป็นการทรมาน การทำที่โหดร้าย ไร้มนุษยธรรม หรือย่ำยีศักดิ์ศรีความเป็นมนุษย์"
                          "หรือกระทำให้บุคคลสูญหายแต่อย่างใด")
        
        doc.add_paragraph(f"เจ้าหน้าที่ผู้จับกุม ได้แจ้งข้อมูลเกี่ยวกับผู้ถูกควบคุมตัว ตามมาตรา ๒๒ วรรคสอง แห่ง พ.ร.บ.ป้องกันและปราบปราม"
                          f"การทรมานและการกระทำให้สูญหาย พ.ศ.๒๕๖๕ ไปยัง ศูนย์ป้องกันปราบปรามทรมานและการทำให้บุคคลสูญหาย"
                          f"ประจำสำนักงานอัยการจังหวัดภูเก็ต เมื่อวันที่ {notify_date} เวลา {notify_attorney} น. เรียบร้อยแล้ว")
        
        doc.add_paragraph(f"เจ้าหน้าที่ผู้จับกุม ได้แจ้งข้อมูลเกี่ยวกับผู้ถูกควบคุมตัว ตามมาตรา ๒๒ วรรคสอง แห่ง พ.ร.บ.ป้องกันและปราบปราม"
                          f"การทรมานและการกระทำให้สูญหาย พ.ศ.๒๕๖๕ ไปยังนายอำเภอเมือง จังหวัดภูเก็ตเป็นผู้รับแจ้งการควบคุมตัว "
                          f"เมื่อวันที่ {notify_date} เวลา {notify_district} น. เรียบร้อยแล้ว\n")
        
        doc.add_paragraph("รับรองว่าข้อความตามบันทึกการจับกุมนี้ถูกต้องตามความเป็นจริงทุกประการ จึงให้ลงลายมือชื่อไว้เป็นหลักฐาน\n\n")
        
        # --- ช่องลายเซ็นผู้ต้องหา (สร้างตามจำนวนคน) ---
        for s in st.session_state.suspects:
            if s['name']:
                p_sign1 = doc.add_paragraph()
                p_sign1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_sign1.add_run(f"(ลงชื่อ)........................................................................ ผู้ถูกจับ/รับสำเนาบันทึกการจับ\n({s['name']})\n\n")
        
        # --- ช่องลายเซ็นชุดจับกุม ---
        raw_officers = officers.replace('\n', ',').replace('และ', ',').split(',')
        officer_list = [o.strip() for o in raw_officers if o.strip()]
        if len(officer_list) > 0:
            p_head = doc.add_paragraph()
            p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_head.add_run(f"(ลงชื่อ)........................................................................ หัวหน้าชุดจับกุม\n({officer_list[0]})\n\n")
            for off in officer_list[1:]:
                p_co = doc.add_paragraph()
                p_co.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_co.add_run(f"(ลงชื่อ)........................................................................ ร่วมจับกุม\n({off})\n\n")
        
        p_reader = doc.add_paragraph()
        p_reader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_reader.add_run("(ลงชื่อ)........................................................................ ผู้บันทึก/อ่าน\n")
        
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    word_file = create_word_doc()
    st.download_button(
        label="ดาวน์โหลดเอกสาร (Export to Word .docx)",
        data=word_file,
        file_name=f"บันทึกจับกุม.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True
    )

# ==========================================
# 📋 กล่องแสดงตัวอย่าง (Preview)
# ==========================================
if st.session_state.show_preview:
    evidence_html = evidence_text_ai.replace('\n', '<br>')
    if not evidence_html:
        evidence_html = "- ไม่มีของกลาง -"
        
    suspect_html = suspect_text_ai.replace('\n', '<br>')
    if not suspect_html:
        suspect_html = "- ยังไม่ได้กรอกข้อมูลผู้ต้องหา -"
        
    behavior_html = behavior_input.replace('\n', '<br>')
    
    st.markdown("---")
    st.markdown("### 📄 ร่างตัวอย่างบันทึกการจับกุม")
    preview_text = f"""
<div style="background-color: #FFFFFF; padding: 30px; border: 1px solid #D6DCE5; border-radius: 4px; color: #333333; font-size: 16px; font-family: 'Sarabun', sans-serif; box-shadow: 0 4px 8px rgba(0,0,0,0.05);">
    <h3 style="text-align: center; color: #0F2C59;">บันทึกการจับกุม</h3>
    <p><b>สถานที่ทำบันทึก:</b> {report_loc}</p>
    <p><b>วัน/เดือน/ปี ที่บันทึก:</b> {report_date}</p>
    <p><b>วัน/เดือน/ปี ที่จับกุม:</b> {arrest_date}</p>
    <p><b>สถานที่เกิดเหตุ/จับกุม:</b> {arrest_loc}</p>
    <br>
    <p><b>นามเจ้าพนักงานผู้จับภายใต้การอำนวยการของ</b> {commanders}</p>
    <p><b>เจ้าหน้าที่ตำรวจชุดจับกุม ได้แก่</b> {officers}</p>
    <br>
    <p><b>ได้ร่วมกันจับกุมตัวผู้ต้องหา คือ</b><br>{suspect_html}</p>
    <br>
    <p><b>พร้อมด้วยของกลาง:</b><br>{evidence_html}</p>
    <br>
    <p><b>โดยกล่าวหาว่า:</b> “{charge_input}”</p>
    <br>
    <p><b>พฤติการณ์ในการจับกุม กล่าวคือ</b><br>{behavior_html}</p>
    <p><b>ในชั้นจับกุม</b> {suspect_statement}</p>
    <br>
    <hr style="border-top: 1px dashed #cccccc;">
    <p style="text-align: center; color: #888888;"><i>(ข้อความแจ้งสิทธิตามกฎหมาย, พ.ร.บ.อุ้มหายฯ และช่องลงนาม จะปรากฏอย่างครบถ้วนในเอกสาร Word ที่ดาวน์โหลด)</i></p>
</div>
    """
    st.markdown(preview_text, unsafe_allow_html=True)
